from selenium import webdriver
from selenium.webdriver.firefox.options import Options
import re
import sys
import traceback
from pathlib import Path
import argparse
import time
import json
import docx


parser = argparse.ArgumentParser()
parser.add_argument("--geckodriver_path",type=str,
    default=str(Path.home()) + "/.local/bin/geckodriver",
    help="Absolute path of the geckodriver binary.")
parser.add_argument("--headless", default=False, action='store_true',
    help="If True: start browser without GUI.")
args = parser.parse_args()

def get_browser_object(geckodriver_path, headless=False):
    # Create browser
    # (Sometimes fails and works again a second later?!?)
    for _ in range(5):
        try:
            options = Options()
            if headless: options.add_argument("--headless")
            return webdriver.Firefox(executable_path=geckodriver_path, options=options)
        except Exception:
            traceback_str = traceback.format_exc()
            time.sleep(2)
    print("Repeatedly failed to create browser object.")
    print("The traceback message is:")
    print(traceback_str)
    print("Exiting script now.")
    sys.exit()

# create folders
Path(Path(__file__).absolute().parent/'visited').mkdir(exist_ok=True)
Path(Path(__file__).absolute().parent/'word_docs').mkdir(exist_ok=True)

with open('search_urls.json', 'r') as f:
    queries = json.load(f)

# Search new job ads for each entry in search_urls.json
for query_name, query_url in list(queries.items()):
    browser = get_browser_object(args.geckodriver_path, args.headless)
    browser.get(query_url)
    browser.implicitly_wait(10)
    try:
        cookies = browser.find_element_by_id('ccmgt_explicit_accept')
        cookies.click()
        browser.implicitly_wait(10)
    except:
        pass

    results = {}
    page = 1
    while True:
        html = browser.page_source
        # Cut out ResultsContainers that each represents a job ad
        job_containers = re.findall('(<article id="job-item)(.*?)(<\/article)', html)
        for _job_html in job_containers:
            job_html = _job_html[1]
            timestamp = re.findall('datetime=(")(.*?)" timeago',
                                   job_html)[0][1][:10]
            _url = ''.join(re.findall('(job-item-title)(.*?)(.html)',
                                      job_html)[0][1:])[8:]
            job_url = 'https://www.stepstone.de/'+_url
            job_title = re.findall('(<h2)(.*?)(<\/h2)',
                                   job_html)[0][1].split('>')[1]
            job_description = re.findall('(<span>)(.*?)(<\/span)', job_html)[0][1]\
                .replace('<strong>','').replace('</strong>','')
            company = re.findall('(<div data-at="job-item-company-name")(.*?)(<\/div)',
                                 job_html)[0][1].split('>')[1]
            location = re.findall('(job-item-location)(.*?)(<\/li)',
                                  job_html)[0][1].split('>')[1]
            results[job_url] = {'timestamp': timestamp,
                                  'job_url': job_url,
                                  'job_title': job_title,
                                  'job_description': job_description,
                                  'company': company,
                                  'location': location}

        try:
            next_page = browser.find_element_by_link_text(str(page+1))
            next_page.click()
            browser.implicitly_wait(10)
            page += 1
        except:
            break
    browser.quit()

    # Keep track of job ads that were aleady added to the word document
    job_json = Path(__file__).absolute().parent/'visited'/f'{query_name}.json'
    if not job_json.exists():
        with(open(job_json, 'w')):
            pass
        visited = results
        new_offers = results
    elif job_json.exists():
        with open(job_json, 'r') as file:
            visited = json.load(file)
        new_offers = {}
        for job_url in results.keys():
            if job_url not in visited.keys():
                visited[job_url] = results[job_url]
                new_offers[job_url] = results[job_url]
    with open(job_json, 'w') as f:
        f.truncate()
        json.dump(visited, f, indent=4)

    # Save job ads to word document
    doc_path = Path(__file__).absolute().parent/'word_docs'/f'{query_name}.docx'
    if not doc_path.exists():
        doc = docx.Document()
    elif doc_path.exists():
        doc = docx.Document(doc_path)
    print(f"Found {len(list(new_offers.keys()))} new jobs for the query {query_name}.")
    for v in new_offers.values():
        doc.add_heading(f"{v['job_title']}", 2)
        doc.add_paragraph(f"Erstellt am: {v['timestamp']}\n")
        doc.add_paragraph(f"Unternehmen: {v['company']}\n")
        doc.add_paragraph(f"Ort: {v['location']}\n")
        doc.add_paragraph(f"{v['job_description']}\n")
        doc.add_paragraph(f"{v['job_url']}")
        doc.add_paragraph("\n")
    doc.add_paragraph("\n")
    doc.save(doc_path)
